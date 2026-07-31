"""
生成 ATS 优化简历 — 测试开发工程师（严格一页A4）
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_line(paragraph):
    """段落下方灰色分隔线"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BBBBBB')
    pBdr.append(bottom)
    pPr.append(pBdr)


def sp(paragraph, before=0, after=0, line_spacing=1.0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def run_font(run, size, name='微软雅黑', bold=False, color=0x333333):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.color.rgb = RGBColor((color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF)


def header(doc, text):
    p = doc.add_paragraph()
    sp(p, before=3, after=0, line_spacing=1.0)
    r = p.add_run(text)
    run_font(r, 9.5, bold=True, color=0x2a2a2a)
    add_line(p)

def proj_title(doc, title, time_range):
    p = doc.add_paragraph()
    sp(p, before=2, after=0, line_spacing=1.0)
    r1 = p.add_run(title)
    run_font(r1, 9, bold=True, color=0x1a1a1a)
    r2 = p.add_run('    ' + time_range)
    run_font(r2, 8, color=0x999999)

def tech_line(doc, text):
    p = doc.add_paragraph()
    sp(p, before=0, after=0, line_spacing=1.0)
    r = p.add_run(text)
    run_font(r, 7.5, color=0x888888)

def bullet(doc, text):
    p = doc.add_paragraph()
    sp(p, before=0, after=0, line_spacing=1.0)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.first_line_indent = Cm(-0.15)
    r = p.add_run('• ' + text)
    run_font(r, 8, color=0x3a3a3a)

def body_line(doc, text, size=8.5, bold=False, color=0x333333, align=None):
    p = doc.add_paragraph()
    sp(p, before=0, after=0, line_spacing=1.0)
    if align:
        p.alignment = align
    r = p.add_run(text)
    run_font(r, size, bold=bold, color=color)
    return p


def build_resume():
    doc = Document()

    # ── 页面：A4 + 最小边距 ──
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(0.5)
    sec.bottom_margin = Cm(0.4)
    sec.left_margin = Cm(1.3)
    sec.right_margin = Cm(1.3)

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(8)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # ========== 个人信息 ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, before=0, after=0, line_spacing=1.0)
    r = p.add_run('高志卫')
    run_font(r, 14, bold=True, color=0x1a1a1a)

    body_line(doc, '19892287896  |  2177685553@qq.com  |  本科 · 计算机科学与技术  |  张家界学院  |  2023.09 - 2027.09',
              size=8, color=0x666666, align=WD_ALIGN_PARAGRAPH.CENTER)
    body_line(doc, '求职方向：测试开发工程师 / 自动化测试工程师  |  意向城市：广州',
              size=7.5, color=0x888888, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ========== 教育经历 ==========
    header(doc, '教育经历')
    body_line(doc, '张家界学院  |  计算机科学与技术（本科）  |  2023.09 - 2027.09', size=8.5)
    body_line(doc, '主修：软件测试基础、数据库原理、计算机网络、操作系统、软件工程、Java程序设计、Python编程、数据结构',
              size=7.5, color=0x777777)

    # ========== 专业技能 ==========
    header(doc, '专业技能')

    skills = [
        ('测试开发', '熟悉软件测试全流程（需求分析→测试设计→缺陷管理→回归测试），掌握等价类划分、边界值分析、场景法、因果图等用例设计方法'),
        ('自动化测试', '熟练Python测试脚本开发，掌握pytest/unittest框架；熟练Selenium Web自动化，掌握Page Object设计模式；熟悉Postman、JMeter接口测试与性能压测'),
        ('AI测试方向', '掌握大模型应用开发技术栈（LangChain、DeepSeek API、Qwen API），熟悉Agent设计模式（ReAct、Plan-Execute、Reflection）；掌握Prompt工程、Function Calling、RAG技术；具备AI驱动测试用例生成与智能测试数据生成的实战经验'),
        ('后端与数据库', '掌握Java + Spring Boot + MyBatis-Plus企业级开发；熟悉MySQL（SQL编写、索引优化、执行计划分析）；了解RESTful API与JWT认证'),
        ('工程化', '熟悉Git分支管理、Linux常用命令、Docker容器化部署；了解Jenkins持续集成与CI/CD流水线'),
    ]

    for cat, detail in skills:
        p = doc.add_paragraph()
        sp(p, before=0, after=0, line_spacing=1.0)
        r1 = p.add_run(f'【{cat}】')
        run_font(r1, 8, bold=True, color=0x2a5da8)
        r2 = p.add_run(f' {detail}')
        run_font(r2, 7.5, color=0x4a4a4a)

    # ========== 项目经历 ==========
    header(doc, '项目经历')

    # ── 项目一 ──
    proj_title(doc, '基于大语言模型Agent的智能测试数据生成与隐私脱敏平台', '2026.03 - 2026.06')
    tech_line(doc, 'Vue3 + Element Plus · Spring Boot 3.3 · MyBatis-Plus · FastAPI · LangChain · DeepSeek / Qwen · MySQL · WebSocket · Faker · TestContainers · Docker')

    for t in [
        '针对测试数据准备成本高、真实数据隐私泄露两大痛点，设计并实现基于LLM Agent的智能测试数据生成与隐私脱敏一体化平台，完整覆盖11个开发阶段',
        '设计三层微服务架构：Vue3前端 + Spring Boot业务编排 + Python FastAPI AI服务，支持DeepSeek/Qwen双模型热切换与LLM失败自动降级为规则引擎',
        '基于LangChain构建多Agent协作系统（SchemaAgent语义理解 + StrategyAgent策略规划 + ToolAgent工具调用），实现ReAct + Plan-Execute混合Agent模式',
        '设计三层递进式敏感数据检测引擎：L1关键词匹配（6类敏感字段映射）→ L2正则检测（手机号/身份证/银行卡/邮箱）→ L3 LLM语义判断，多层融合实现高召回低误报',
        '基于JDBC读取information_schema元数据，自动解析表结构、字段类型、主外键关系，构建Kahn拓扑排序算法解析多表依赖并检测循环依赖',
        '实现10种数据生成器（Faker真实模拟 + 数值/时间/枚举/常量），通过GenerationContext维护外键关联一致性，支持单表/多表批量生成与参数化SQL写入',
        '编写95个单元测试全部通过，基于TestContainers实现12步全流程集成测试（Schema解析→LLM规划→数据生成→FK校验→敏感检测→脱敏→CSV/SQL/JSON导出）',
    ]:
        bullet(doc, t)

    # ── 项目二 ──
    proj_title(doc, 'AI智能自动化测试平台', '2025.10 - 2025.12')
    tech_line(doc, 'React 18 + TypeScript · FastAPI · PostgreSQL · Redis · OpenAI GPT-4 / DeepSeek · Ollama · Selenium · k6 · Docker Compose · WebSocket')

    for t in [
        '独立设计并实现AI驱动的全栈自动化测试平台，覆盖需求智能分析、测试用例生成、API测试、UI自动化、性能测试五大场景，支持从需求到报告的全流程自动化闭环',
        '基于LLM实现自然语言需求智能分析，自动提取测试要点、边界条件及风险点；支持OpenAPI/Swagger/Postman文档解析，AI自动生成接口测试脚本（正常/异常/边界场景）',
        '实现AI驱动的UI测试脚本生成：基于页面DOM结构智能识别元素定位策略（类Testim自愈式定位），自动生成Selenium/Playwright测试脚本',
        '集成k6性能测试引擎，AI自动解析测试结果（响应时间/吞吐量/错误率）并生成优化建议；支持异步并行执行与WebSocket实时进度推送',
        '采用工厂+策略模式实现多AI模型（GPT-4/DeepSeek/Ollama）统一接入与热切换，支持Docker Compose一键部署',
    ]:
        bullet(doc, t)

    # ── 项目三 ──
    proj_title(doc, '基于Selenium + Page Object的Web自动化测试框架', '2025.07 - 2025.09')
    tech_line(doc, 'Python · Selenium WebDriver · Pytest · Page Object Model · Allure · Jenkins')

    for t in [
        '面向Web系统回归测试效率低的痛点，设计基于Page Object模式的自动化测试框架，实现页面操作与测试逻辑分离，提升脚本可维护性60%+',
        '使用Selenium WebDriver实现Web端到端自动化操作（登录/查询/表单/业务流程），封装BasePage基类统一管理元素等待、截图、日志等公共行为',
        '基于pytest框架组织测试用例并实现参数化执行，集成conftest.py进行全局fixture管理（driver生命周期、测试数据准备、环境配置）',
        '实现失败自动截图 + 日志记录 + Allure可视化报告三重定位机制，集成Jenkins实现定时回归执行',
    ]:
        bullet(doc, t)

    # ========== 个人评价 ==========
    header(doc, '个人评价')
    body_line(doc,
        '计算机科学与技术专业本科生，具备扎实的软件测试理论基础、自动化测试开发及后端开发能力。'
        '熟悉Web端与接口测试全流程，能独立使用Python搭建测试框架，具备Java企业级项目开发经验。'
        '深度关注AI技术在测试领域的应用落地，具备LangChain Agent开发、LLM应用集成和Prompt工程实践经验。'
        '学习能力强，善于问题分析与技术攻坚，具备独立完成毕业设计项目（11阶段92%完成度）的工程能力，希望在测试开发方向持续深耕。',
        size=7.5, color=0x4a4a4a)

    # ========== 保存 ==========
    out = r'C:\Users\21776\OneDrive\桌面\高志卫_测试开发工程师_简历.docx'
    doc.save(out)
    print(f'[OK] {out}')
    return out


if __name__ == '__main__':
    build_resume()
