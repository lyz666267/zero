"""
生成简历 v3 — 不限一页 · 突出职责与量化指标 · ATS优化
基于模板: 微信文件 简历.docx (2026-08) 视觉风格
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 颜色常量 ──
COLOR_DEEP_BLUE = 0x1B3A5C      # 标题/章节头/边框/bullet
COLOR_PROJECT_TITLE = 0x2B579A  # 项目标题
COLOR_META = 0x666666           # 元信息（时间/技术栈）
COLOR_BODY = 0x333333           # 正文
COLOR_ROLE_BG = 'E8F0F8'       # 角色标签背景色（浅蓝）


def add_bottom_border(paragraph, color_hex='1B3A5C', sz='4', space='4'):
    """底部边框 — 模板风格深蓝细线"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), space)
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def sp(paragraph, before=0, after=0, line=1.0):
    """段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def rf(run, size_pt, bold=False, color_rgb=None, font_name='微软雅黑'):
    """设置 run 字体"""
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color_rgb is not None:
        run.font.color.rgb = RGBColor(
            (color_rgb >> 16) & 0xFF,
            (color_rgb >> 8) & 0xFF,
            color_rgb & 0xFF
        )


def ar(paragraph, text, size_pt, bold=False, color_rgb=COLOR_BODY):
    """添加带格式的 run"""
    r = paragraph.add_run(text)
    rf(r, size_pt, bold, color_rgb)
    return r


# ==================== 组件函数 ====================

def title_line(doc, text):
    """大标题 — 16pt bold #1B3A5C 居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, before=0, after=4, line=1.0)
    ar(p, text, 14, bold=True, color_rgb=COLOR_DEEP_BLUE)


def section_header(doc, text):
    """章节标题 — 11pt bold #1B3A5C + 底部蓝色细线"""
    p = doc.add_paragraph()
    sp(p, before=4, after=1, line=1.0)
    ar(p, text, 10, bold=True, color_rgb=COLOR_DEEP_BLUE)
    add_bottom_border(p, '1B3A5C', '4', '4')


def info_line(doc, label, value, size=9):
    """信息行 — bold标签 + 普通值"""
    p = doc.add_paragraph()
    sp(p, before=0, after=0, line=1.05)
    ar(p, label, size, bold=True)
    ar(p, value, size, bold=False)


def skill_block(doc, label, detail, size=8.5):
    """技能块 — bold类别标签 + 内容"""
    p = doc.add_paragraph()
    sp(p, before=0, after=0, line=1.1)
    ar(p, label, size, bold=True)
    ar(p, detail, size, bold=False)


def project_title_line(doc, text):
    """项目标题 — 10pt bold #2B579A"""
    p = doc.add_paragraph()
    sp(p, before=4, after=0, line=1.0)
    ar(p, text, 10, bold=True, color_rgb=COLOR_PROJECT_TITLE)


def project_meta_line(doc, text):
    """项目元信息 — 8pt #666666（时间 + 技术栈 + 概述）"""
    p = doc.add_paragraph()
    sp(p, before=0, after=1, line=1.05)
    ar(p, text, 8, bold=False, color_rgb=COLOR_META)


def role_line(doc, text):
    """个人职责 — 8.5pt bold深蓝标签 + 正文"""
    p = doc.add_paragraph()
    sp(p, before=0, after=1, line=1.1)
    p.paragraph_format.left_indent = Cm(0.3)
    ar(p, '▎我的职责：', 8.5, bold=True, color_rgb=COLOR_DEEP_BLUE)
    ar(p, text, 8.5, bold=False)


def bullet_point(doc, text, size=8.5):
    """项目要点 — 深蓝bullet + 正文，0.5cm缩进"""
    p = doc.add_paragraph()
    sp(p, before=0, after=0, line=1.1)
    p.paragraph_format.left_indent = Cm(0.5)
    ar(p, '• ', size, bold=False, color_rgb=COLOR_DEEP_BLUE)
    ar(p, text, size, bold=False)


def metric_bullet(doc, metric, text, size=8.5):
    """带量化前置标签的要点 — 【指标】正文"""
    p = doc.add_paragraph()
    sp(p, before=0, after=0, line=1.1)
    p.paragraph_format.left_indent = Cm(0.5)
    ar(p, '• ', size, bold=False, color_rgb=COLOR_DEEP_BLUE)
    ar(p, f'{metric}：', size, bold=True)
    ar(p, text, size, bold=False)


def campus_bullet(doc, label, text, size=8.5):
    """校园经历/自我评价要点"""
    p = doc.add_paragraph()
    sp(p, before=0, after=0, line=1.1)
    p.paragraph_format.left_indent = Cm(0.5)
    ar(p, '• ', size, bold=False, color_rgb=COLOR_DEEP_BLUE)
    ar(p, label, size, bold=True)
    ar(p, text, size, bold=False)


# ==================== 构建简历 ====================

def build_resume():
    doc = Document()

    # ── 页面设置（A4，舒适边距）──
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.2)
    sec.bottom_margin = Cm(1.0)
    sec.left_margin = Cm(1.6)
    sec.right_margin = Cm(1.6)

    # ── 默认样式 ──
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(8.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # ==================== 标题 ====================
    title_line(doc, '个人简历')

    # ==================== 基本信息 ====================
    section_header(doc, '基本信息')
    info_line(doc, '姓　　名：', '高志卫')
    info_line(doc, '出生年月：', '2005年6月')
    info_line(doc, '毕业院校：', '张家界学院（2023.09 - 2027.06）')
    info_line(doc, '专　　业：', '计算机科学与技术（本科）')
    info_line(doc, '求职意向：', '测试开发工程师 / 自动化测试工程师（广州）')
    info_line(doc, '手机号码：', '19892287896')
    info_line(doc, '电子邮箱：', '2177685553@qq.com')

    # ==================== 专业技能 ====================
    section_header(doc, '专业技能')

    skill_block(doc,
        '测试理论与方法：',
        '掌握软件测试全流程（需求分析→测试设计→缺陷管理→回归测试），'
        '熟练运用等价类划分、边界值分析、场景法、因果图、判定表、正交试验等用例设计方法；'
        '熟悉测试计划制定、测试报告编写、缺陷生命周期管理（JIRA/禅道）')

    skill_block(doc,
        '自动化测试：',
        '精通Python测试脚本开发，熟练使用pytest/unittest框架进行用例组织与参数化执行；'
        '熟练Selenium WebDriver实现Web端到端自动化，掌握Page Object设计模式；'
        '熟悉Postman、JMeter进行接口测试与性能压测，能独立构建自动化测试框架')

    skill_block(doc,
        'AI测试技术：',
        '掌握大模型应用开发技术栈（LangChain、DeepSeek API、Qwen API），'
        '熟悉Agent设计模式（ReAct、Plan-Execute、Reflection、Tool-Use）；'
        '掌握Prompt工程、Function Calling、RAG技术；'
        '具备AI驱动测试用例生成与智能测试数据生成的完整项目实战经验')

    skill_block(doc,
        '后端与数据库：',
        '掌握Java + Spring Boot 3 + MyBatis-Plus企业级开发，熟悉RESTful API设计、'
        'JWT认证授权、全局异常处理、AOP日志切面等后端通用能力；'
        '熟悉MySQL数据库（SQL编写、索引优化、EXPLAIN执行计划分析），了解Redis缓存策略')

    skill_block(doc,
        '工程化与工具链：',
        '熟练Git分支管理、Linux常用命令、Docker容器化部署与Docker Compose服务编排；'
        '了解Jenkins持续集成与CI/CD流水线配置；'
        '熟悉HTTP协议请求/响应模型、常见状态码含义及RESTful规范')

    # ==================== 项目经历 ====================
    section_header(doc, '项目经历')

    # ========== 项目一 ==========
    project_title_line(doc, '项目一：基于大语言模型Agent的智能测试数据生成与隐私脱敏平台')
    project_meta_line(doc,
        '2026.03 - 2026.06  |  Vue3 + Spring Boot 3 + FastAPI + LangChain + DeepSeek/Qwen + MySQL + TestContainers + Docker  |  '
        '毕业设计课题，针对测试数据准备成本高、真实数据隐私泄露两大行业痛点，'
        '独立设计并实现基于LLM Agent的智能测试数据生成与隐私脱敏一体化平台')

    role_line(doc,
        '作为项目负责人独立完成全周期开发（11个阶段），负责系统架构设计、后端业务编排、AI服务开发、'
        '敏感数据检测引擎设计、数据生成器开发、集成测试编写等核心工作，项目整体完成度达92%')

    metric_bullet(doc, '系统架构',
        '设计三层微服务架构（Vue3前端 → Spring Boot业务编排 → Python FastAPI AI服务），'
        '前后端分离部署，RESTful API + JSON通信；支持DeepSeek/Qwen双模型热切换与LLM调用失败自动降级为规则引擎（Mock模式），'
        '保障测试数据生成服务的高可用')

    metric_bullet(doc, 'Agent协作系统',
        '基于LangChain构建多Agent协作系统：SchemaAgent（数据库Schema语义理解与字段类型推断）→ '
        'StrategyAgent（基于ReAct + Plan-Execute模式制定生成策略）→ ToolAgent（调用10种数据生成器执行生成），'
        '实现从Schema解析到数据输出的全自动编排')

    metric_bullet(doc, '敏感数据检测',
        '设计三层递进式敏感数据检测引擎：L1层关键词匹配（覆盖姓名/手机/身份证/银行卡/邮箱/地址6类敏感字段）→ '
        'L2层正则模式检测（手机号/身份证号/银行卡号/邮箱4种正则规则）→ L3层LLM语义判断，'
        '多层融合策略兼顾检测召回率与误报率')

    metric_bullet(doc, '依赖解析与拓扑排序',
        '基于JDBC读取MySQL information_schema元数据，自动解析表结构（字段名/类型/长度/是否可空）、主键、外键关系；'
        '实现Kahn拓扑排序算法构建表间依赖关系图，自动推导多表生成顺序并检测循环依赖；'
        '通过GenerationContext维护外键关联一致性，确保多表联生成时引用关系完整')

    metric_bullet(doc, '数据生成引擎',
        '实现10种数据生成器（Faker真实模拟/随机数值/时间序列/枚举选择/常量填充/正则匹配/外键引用/'
        '自增序列/表达式计算/级联生成），覆盖116个字段语义→生成器映射；'
        '支持单表/多表批量生成，参数化配置生成行数（1~10000行），批量SQL INSERT写入')

    metric_bullet(doc, '测试质量保障',
        '编写95个单元测试（pytest），覆盖字段生成器/类型映射/敏感检测/拓扑排序等核心模块，全部通过；'
        '基于TestContainers启动真实MySQL容器，设计12步全流程集成测试'
        '（Schema解析→LLM规划→数据生成→FK校验→敏感检测→脱敏处理→CSV/SQL/JSON三格式导出），端到端验证闭环')

    # ========== 项目二 ==========
    project_title_line(doc, '项目二：AI智能自动化测试平台')
    project_meta_line(doc,
        '2025.10 - 2025.12  |  React + TypeScript + FastAPI + PostgreSQL + Redis + GPT-4/DeepSeek/Ollama + Selenium + k6 + Docker  |  '
        '独立设计并实现AI驱动的全栈自动化测试平台，覆盖从需求分析到测试报告输出的全流程自动化闭环')

    role_line(doc,
        '独立完成平台全栈开发，负责需求分析模块、AI接口测试生成、UI自动化脚本生成、k6性能测试集成、'
        '多模型调度框架等核心功能的设计与实现')

    metric_bullet(doc, '需求智能分析',
        '基于LLM实现自然语言需求文档自动解析，提取测试要点、边界条件及潜在风险点；'
        '支持OpenAPI/Swagger/Postman Collection三种接口文档导入，AI自动生成接口测试脚本'
        '（覆盖正常场景/异常场景/边界场景3类用例），将手工编写测试脚本时间缩短约60%')

    metric_bullet(doc, 'UI自动化测试',
        '实现AI驱动的UI测试脚本自动生成：解析页面DOM结构，智能识别元素定位策略'
        '（优先ID→Name→CSS Selector→XPath，类Testim自愈式定位），自动生成Selenium/Playwright可执行测试脚本；'
        '支持登录/表单/搜索/列表/弹窗5类常见UI场景的元素识别与操作封装')

    metric_bullet(doc, '性能测试集成',
        '集成k6性能测试引擎，AI自动解析测试结果关键指标（响应时间P50/P95/P99、吞吐量RPS、错误率），'
        '结合阈值规则自动生成性能优化建议；支持并发用户数参数化配置（10~1000 VUs）与异步并行执行，'
        '通过WebSocket实时推送测试进度到前端')

    metric_bullet(doc, '多模型调度',
        '采用工厂模式 + 策略模式实现多AI模型（GPT-4/DeepSeek/Ollama本地模型）统一接口接入与运行时热切换；'
        '支持Docker Compose一键部署（React前端 + FastAPI后端 + PostgreSQL + Redis），启动时间约30秒')

    # ========== 项目三 ==========
    project_title_line(doc, '项目三：基于Selenium + Page Object的Web自动化测试框架')
    project_meta_line(doc,
        '2025.07 - 2025.09  |  Python · Selenium WebDriver · Pytest · Page Object Model · Allure · Jenkins  |  '
        '面向Web系统回归测试效率低的痛点，设计基于Page Object模式的自动化测试框架，集成CI/CD流水线实现定时回归执行')

    role_line(doc,
        '负责框架整体架构设计、BasePage基类封装、测试用例编写与参数化、Jenkins CI/CD流水线配置及Allure报告集成')

    metric_bullet(doc, '框架设计',
        '基于Page Object设计模式将页面元素定位与业务操作封装为独立Page类（LoginPage/SearchPage/FormPage等），'
        '实现页面操作与测试逻辑分离；封装BasePage基类统一管理显式等待（WebDriverWait）、失败自动截图、'
        '日志记录等公共行为，脚本可维护性提升约60%')

    metric_bullet(doc, '用例组织与参数化',
        '基于pytest框架组织30+条测试用例，通过@pytest.mark.parametrize实现多组数据驱动执行；'
        '集成conftest.py进行全局fixture管理（driver生命周期管理、测试数据准备与清理、多环境配置切换），'
        '单次回归执行覆盖登录/查询/新增/编辑/删除/批量操作6类核心业务流程')

    metric_bullet(doc, 'CI/CD与报告',
        '实现失败自动截图 + 结构化日志 + Allure可视化报告三重定位机制，测试结果可追溯；'
        '配置Jenkins Job实现每日凌晨4点定时回归执行，测试报告自动归档并通过邮件通知相关人员')

    # ==================== 校园经历 ====================
    section_header(doc, '校园经历')

    campus_bullet(doc, '课程项目实践：',
        '在「软件测试」「Java Web程序设计」「数据库原理」等课程中担任小组负责人，'
        '协调5-6人团队完成从需求分析到代码交付的全过程，负责任务拆分、进度跟进、代码Review与最终答辩汇报，'
        '多个项目获课程优秀评级')

    campus_bullet(doc, '毕业设计：',
        '独立完成「基于大语言模型Agent的智能测试数据生成与隐私脱敏平台」毕业设计课题，'
        '完整经历需求分析→技术选型→架构设计→编码实现→测试验证→文档撰写全周期，'
        '覆盖11个开发阶段，完成度达92%，编写95个单元测试及12步全流程集成测试')

    # ==================== 自我评价 ====================
    section_header(doc, '自我评价')

    campus_bullet(doc, '测试专业能力：',
        '计算机科班背景，系统学习软件测试、数据库原理、计算机网络、操作系统等核心课程；'
        '掌握黑盒测试方法及多种用例设计技术，在3个项目中积累了自动化测试框架搭建、测试用例编写、'
        '缺陷管理与回归测试的实战经验')

    campus_bullet(doc, 'AI测试实践：',
        '深度关注AI技术在测试领域的应用落地，具备LangChain Agent开发、LLM应用集成、'
        'Prompt工程和RAG技术的完整实践经验；能结合实际测试场景（用例生成、数据构造、结果分析）进行AI赋能探索，'
        '有2个AI驱动测试工具从0到1的独立开发经历')

    campus_bullet(doc, '工程素养与协作：',
        '测试执行时注重截图留证，缺陷提单时附完整复现步骤与环境信息；'
        '在课程项目中多次担任小组组长，善于将复杂任务拆解为可执行的子任务并设定里程碑，'
        '能通过清晰的技术文档和定期沟通与团队成员高效协作')

    # ── 保存 ──
    out = r'C:\Users\21776\OneDrive\桌面\高志卫_测试开发工程师_简历.docx'
    doc.save(out)
    print(f'[OK] DOCX saved: {out}')
    return out


if __name__ == '__main__':
    build_resume()
