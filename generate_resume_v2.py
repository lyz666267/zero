"""
生成简历 — 复刻唐宇佳模板视觉风格（压缩版适配1-2页A4）
基于模板: 微信文件 简历.docx (2026-08)

模板风格保留:
- 标题/章节头 #1B3A5C 深蓝色 + 底部蓝色细线
- 项目标题 #2B579A 蓝色
- 元信息 #666666 灰色
- bullet符 #1B3A5C 色
- 信息行 bold标签 + 普通值
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# 颜色常量 (与模板完全一致)
COLOR_DEEP_BLUE = 0x1B3A5C
COLOR_PROJECT_TITLE = 0x2B579A
COLOR_META = 0x666666


def add_bottom_border(paragraph, color_hex='1B3A5C', sz='4', space='4'):
    """底部边框 — 模板风格"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), space)
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def sp(paragraph, before=0, after=0, line=1.0):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def rf(run, size_pt, bold=False, color_rgb=None, font_name='微软雅黑'):
    """设置 run 字体"""
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color_rgb is not None:
        run.font.color.rgb = RGBColor(
            (color_rgb >> 16) & 0xFF,
            (color_rgb >> 8) & 0xFF,
            color_rgb & 0xFF
        )


def ar(paragraph, text, size_pt, bold=False, color_rgb=None):
    """添加带格式的 run"""
    r = paragraph.add_run(text)
    rf(r, size_pt, bold, color_rgb)
    return r


# ==================== 组件函数 ====================

def title_line(doc, text):
    """大标题 — 14pt bold #1B3A5C 居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, after=0, line=1.0)
    ar(p, text, 12, bold=True, color_rgb=COLOR_DEEP_BLUE)


def section_header(doc, text):
    """章节标题 — 10pt bold #1B3A5C + 底部蓝色细线"""
    p = doc.add_paragraph()
    sp(p, before=0, after=0, line=1.0)
    ar(p, text, 9, bold=True, color_rgb=COLOR_DEEP_BLUE)
    add_bottom_border(p, '1B3A5C', '4', '4')


def info_line(doc, label, value, size=7):
    """信息行 — bold标签 + 普通值"""
    p = doc.add_paragraph()
    sp(p, line=1.0)
    ar(p, label, size, bold=True)
    ar(p, value, size, bold=False)


def skill_line(doc, label, detail, size=7):
    """技能行 — bold标签 + 普通内容"""
    p = doc.add_paragraph()
    sp(p, line=1.0)
    ar(p, label, size, bold=True)
    ar(p, detail, size, bold=False)


def project_title_line(doc, text):
    """项目标题 — 9pt bold #2B579A"""
    p = doc.add_paragraph()
    sp(p, before=1, line=1.0)
    ar(p, text, 9, bold=True, color_rgb=COLOR_PROJECT_TITLE)


def project_meta_line(doc, text):
    """项目元信息 — 7.5pt #666666"""
    p = doc.add_paragraph()
    sp(p, line=1.0)
    ar(p, text, 7.5, bold=False, color_rgb=COLOR_META)


def bullet_point(doc, text, size=7):
    """项目要点 — bullet #1B3A5C + 正文 indent"""
    p = doc.add_paragraph()
    sp(p, line=1.0)
    p.paragraph_format.left_indent = Cm(0.4)
    ar(p, '• ', size, bold=False, color_rgb=COLOR_DEEP_BLUE)
    ar(p, text, size, bold=False)


def star_bullet(doc, label, text, size=7):
    """校园经历/自我评价要点 — bold子标签"""
    p = doc.add_paragraph()
    sp(p, line=1.0)
    p.paragraph_format.left_indent = Cm(0.4)
    ar(p, '• ', size, bold=False, color_rgb=COLOR_DEEP_BLUE)
    ar(p, label, size, bold=True)
    ar(p, text, size, bold=False)


# ==================== 构建简历 ====================

def build_resume():
    doc = Document()

    # ── 页面设置 ──
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(0.5)
    sec.bottom_margin = Cm(0.2)
    sec.left_margin = Cm(1.2)
    sec.right_margin = Cm(1.2)

    # ── 默认样式 ──
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(7)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # ==================== 标题 ====================
    title_line(doc, '个人简历')

    # ==================== 基本信息 ====================
    section_header(doc, '基本信息')
    info_line(doc, '姓名：', '高志卫')
    info_line(doc, '出生年月：', '2003年6月')
    info_line(doc, '毕业院校：', '张家界学院（2023.09 - 2027.06）')
    info_line(doc, '专　　业：', '计算机科学与技术（本科）')
    info_line(doc, '求职意向：', '测试开发工程师 / 自动化测试工程师（广州）')
    info_line(doc, '手机号码：', '19892287896')
    info_line(doc, '邮　　箱：', '2177685553@qq.com')

    # ==================== 专业技能 ====================
    section_header(doc, '专业技能')
    skill_line(doc, '测试理论：',
               '掌握黑盒测试方法，熟悉等价类划分、边界值分析、场景法、因果图等用例设计方法，'
               '了解软件测试全流程（需求分析→测试设计→缺陷管理→回归测试）及SDLC')
    skill_line(doc, '自动化测试：',
               '熟练Python测试脚本开发，掌握pytest/unittest框架；熟练Selenium Web自动化，掌握Page Object设计模式；'
               '熟悉Postman、JMeter接口测试与性能压测')
    skill_line(doc, 'AI测试方向：',
               '掌握大模型应用开发技术栈（LangChain、DeepSeek API、Qwen API），熟悉Agent设计模式（ReAct、Plan-Execute、Reflection）；'
               '掌握Prompt工程、Function Calling、RAG技术；具备AI驱动测试用例生成与智能测试数据生成的实战经验')
    skill_line(doc, '后端与数据库：',
               '掌握Java + Spring Boot + MyBatis-Plus企业级开发；熟悉MySQL（SQL编写、索引优化、执行计划分析）；'
               '了解RESTful API、JWT认证与Redis缓存策略')
    skill_line(doc, '其他：',
               '熟悉Git分支管理、Linux常用命令、Docker容器化部署；了解Jenkins持续集成与CI/CD流水线；'
               '了解HTTP协议请求/响应模型及常见状态码含义')

    # ==================== 项目经历 ====================
    section_header(doc, '项目经历')

    # ── 项目一 ──
    project_title_line(doc, '项目一：基于大语言模型Agent的智能测试数据生成与隐私脱敏平台')
    project_meta_line(doc, '2026.03 - 2026.06  |  '
                      'Vue3 + Spring Boot + FastAPI + LangChain + DeepSeek/Qwen + MySQL + TestContainers + Docker  |  '
                      '针对测试数据准备成本高、隐私泄露两大痛点，独立设计并实现基于LLM Agent的智能测试数据生成与隐私脱敏一体化平台，完整覆盖11个开发阶段')
    bullet_point(doc,
        '【架构设计】设计三层微服务架构（Vue3前端 + Spring Boot业务编排 + Python FastAPI AI服务），'
        '支持DeepSeek/Qwen双模型热切换与LLM失败自动降级为规则引擎（Mock模式），保障服务高可用')
    bullet_point(doc,
        '【Agent协作】基于LangChain构建多Agent协作系统（SchemaAgent语义理解 + StrategyAgent策略规划 + '
        'ToolAgent工具调用），实现ReAct + Plan-Execute混合Agent模式')
    bullet_point(doc,
        '【敏感检测】设计三层递进式敏感数据检测引擎：L1关键词匹配（6类敏感字段）→ L2正则检测'
        '（手机号/身份证/银行卡/邮箱）→ L3 LLM语义判断，多层融合实现高召回低误报')
    bullet_point(doc,
        '【依赖解析】基于JDBC读取information_schema自动解析表结构与主外键关系，实现Kahn拓扑排序算法'
        '解析多表依赖并检测循环依赖，通过GenerationContext维护外键关联一致性')
    bullet_point(doc,
        '【数据生成】实现10种数据生成器（Faker真实模拟 + 数值/时间/枚举/常量），'
        '支持单表/多表批量生成与参数化SQL写入；基于TestContainers完成12步全流程集成测试，95个单元测试全部通过')

    # ── 项目二 ──
    project_title_line(doc, '项目二：AI智能自动化测试平台')
    project_meta_line(doc, '2025.10 - 2025.12  |  '
                      'React + TypeScript + FastAPI + PostgreSQL + Redis + GPT-4/DeepSeek/Ollama + Selenium + k6 + Docker  |  '
                      '独立设计并实现AI驱动的全栈自动化测试平台，覆盖需求分析→测试用例生成→API/UI/性能测试→报告输出的全流程自动化闭环')
    bullet_point(doc,
        '【需求分析】基于LLM实现自然语言需求智能分析，自动提取测试要点、边界条件及风险点；'
        '支持OpenAPI/Swagger/Postman文档解析，AI自动生成接口测试脚本（正常/异常/边界场景）')
    bullet_point(doc,
        '【UI自动化】实现AI驱动的UI测试脚本生成：基于页面DOM结构智能识别元素定位策略'
        '（类Testim自愈式定位），自动生成Selenium/Playwright测试脚本')
    bullet_point(doc,
        '【性能测试】集成k6性能测试引擎，AI自动解析测试结果（响应时间/吞吐量/错误率）并生成优化建议；'
        '支持异步并行执行与WebSocket实时进度推送')
    bullet_point(doc,
        '【模型调度】采用工厂+策略模式实现多AI模型统一接入与热切换，支持Docker Compose一键部署')

    # ── 项目三 ──
    project_title_line(doc, '项目三：基于Selenium + Page Object的Web自动化测试框架')
    project_meta_line(doc, '2025.07 - 2025.09  |  '
                      'Python · Selenium WebDriver · Pytest · Page Object Model · Allure · Jenkins  |  '
                      '面向Web系统回归测试效率低的痛点，设计基于Page Object模式的自动化测试框架，实现页面操作与测试逻辑分离，集成CI/CD定时回归执行')
    bullet_point(doc,
        '【框架设计】使用Selenium WebDriver实现Web端到端自动化操作（登录/查询/表单/业务流程），'
        '封装BasePage基类统一管理元素等待、截图、日志等公共行为')
    bullet_point(doc,
        '【用例管理】基于pytest框架组织测试用例并实现参数化执行，集成conftest.py进行全局fixture管理'
        '（driver生命周期、测试数据准备、环境配置）')
    bullet_point(doc,
        '【CI/CD集成】实现失败自动截图 + 日志记录 + Allure可视化报告三重定位机制，集成Jenkins实现定时回归执行')

    # ==================== 校园经历 ====================
    section_header(doc, '校园经历')
    star_bullet(doc, '课程项目实践：',
                '在「软件测试」「Java Web程序设计」等课程中担任小组负责人，'
                '协调5-6人团队完成从需求分析到代码交付的全过程，负责任务拆分、进度跟进与最终汇报')
    star_bullet(doc, '毕业设计：',
                '独立完成「基于大语言模型Agent的智能测试数据生成与隐私脱敏平台」毕业设计课题，'
                '完整经历需求分析→架构设计→编码实现→测试验证全周期，完成度达92%（11阶段）')

    # ==================== 自我评价 ====================
    section_header(doc, '自我评价')
    star_bullet(doc, '专业基础扎实：',
                '计算机科班背景，系统学习过软件测试、数据库原理、计算机网络等核心课程，'
                '掌握黑盒测试方法和用例设计方法，在多个项目中积累了实际的测试用例编写和自动化测试经验')
    star_bullet(doc, 'AI测试实践：',
                '深度关注AI技术在测试领域的应用落地，具备LangChain Agent开发、'
                'LLM应用集成和Prompt工程实践经验，能结合实际测试场景进行AI赋能探索')
    star_bullet(doc, '注重细节与协作：',
                '测试执行时注重截图留证，缺陷提单时附完整的复现步骤；'
                '在课程项目中多次担任小组组长，善于将复杂任务拆解为可执行的子任务，能与团队成员高效分工配合')

    # ========== 保存 ==========
    out = r'C:\Users\21776\OneDrive\桌面\高志卫_测试开发工程师_简历.docx'
    doc.save(out)
    print(f'[OK] DOCX saved: {out}')
    return out


if __name__ == '__main__':
    build_resume()
